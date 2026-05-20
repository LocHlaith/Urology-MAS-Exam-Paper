"""将人类题库 Word 文件转换为 TXT 中间文件。

脚本用途：把人类题库 Word 文档导出为纯文本，供结构化解析脚本读取。
流程阶段：人类题库结构化。
主要输入：用户在文件选择窗口中选定的 `.docx` 或 `.doc` 文件。
主要输出：与 Word 文件同名的 `.txt` 中间文件。
重要边界：人类题库来源是 Word；TXT 只是解析中间文件，不是独立来源，也不是 MAS 出题结果。
"""

import os
import sys
import tkinter as tk
from tkinter import filedialog, messagebox


# ===== Word 转 TXT =====

def convert_docx_to_txt(input_path: str, output_path: str) -> None:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("缺少依赖 python-docx，请先执行：pip install python-docx") from e

    doc = Document(input_path)

    lines = []
    # 段落
    for p in doc.paragraphs:
        text = (p.text or "").rstrip()
        if text != "":
            lines.append(text)

    # 表格
    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = " ".join((cell.text or "").split())
                cells.append(cell_text)
            line = "\t".join(cells).rstrip()
            if line.strip():
                lines.append(line)

    content = "\n".join(lines).strip() + "\n"

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)


def convert_doc_to_txt_windows(input_path: str, output_path: str) -> None:
    if os.name != "nt":
        raise RuntimeError(".doc 转换需要 Windows + Microsoft Word（COM 自动化）。")

    try:
        import win32com.client  # type: ignore
    except ImportError as e:
        raise RuntimeError("缺少依赖 pywin32，请先执行：pip install pywin32") from e

    # Word 常量（避免依赖 win32com.client.constants）
    wdDoNotSaveChanges = 0
    wdFormatText = 2  # 纯文本

    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        # 仅使用必要的 Word Open 参数，并保持只读。
        doc = word.Documents.Open(input_path, ReadOnly=True)

        # 直接让 Word “另存为” TXT，避免自行解析旧版 .doc。
        # 注意：Word 可能会弹编码/兼容性提示，下面两行尽量压制交互
        word.DisplayAlerts = 0

        doc.SaveAs(output_path, FileFormat=wdFormatText)

    finally:
        try:
            if doc is not None:
                doc.Close(SaveChanges=wdDoNotSaveChanges)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass


# ===== 文件选择 =====

def pick_file() -> str:
    root = tk.Tk()
    root.withdraw()
    root.update()

    file_path = filedialog.askopenfilename(
        title="选择 Word 文件（.docx 或 .doc）",
        filetypes=[
            ("Word 文件", "*.docx *.doc"),
            ("DOCX", "*.docx"),
            ("DOC", "*.doc"),
            ("所有文件", "*.*"),
        ],
    )
    root.destroy()
    return file_path


# ===== 程序入口 =====

def main():
    input_path = pick_file()
    if not input_path:
        return

    ext = os.path.splitext(input_path)[1].lower()
    output_path = os.path.splitext(input_path)[0] + ".txt"

    try:
        if ext == ".docx":
            convert_docx_to_txt(input_path, output_path)
        elif ext == ".doc":
            convert_doc_to_txt_windows(input_path, output_path)
        else:
            raise RuntimeError(f"不支持的文件类型：{ext}")

        messagebox.showinfo("转换完成", f"已生成：\n{output_path}")

    except Exception as e:
        messagebox.showerror("转换失败", str(e))
        raise


if __name__ == "__main__":
    main()
