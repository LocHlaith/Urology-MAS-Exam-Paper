"""将 MAS 题库 JSON 导出为人工审阅用 Word 文档。

脚本用途：把 MAS 题库 JSON 转成便于合作者人工审阅的 Word 文档。
流程阶段：人工审阅导出。
主要输入：`data/banks/new_bank_*.json`。
主要输出：`outputs/report_exports/*.docx`。
重要边界：本脚本不产生统计结果，也不是绘图数据来源。
"""

import json
import os
import re
import sys
from pathlib import Path
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
from project_paths import BANK_DIR, REPORT_EXPORTS_DIR


# ===== 路径与常量 =====

TEST_POINTS = {
    1: "尿液检查-血尿", 2: "尿液检查-蛋白尿", 3: "尿液检查-管型尿",
    4: "肾小球疾病-概述", 5: "肾小球疾病-急性肾小球肾炎", 6: "肾小球疾病-急进性肾小球肾炎",
    7: "肾小球疾病-慢性肾小球肾炎", 8: "肾小球疾病-肾病综合征", 9: "肾小球疾病-IgA肾病",
    10: "肾间质疾病-急性间质性肾炎", 11: "尿路感染-概述", 12: "尿路感染-急性肾盂肾炎",
    13: "尿路感染-慢性肾盂肾炎", 14: "尿路感染-急性膀胱炎", 15: "尿路感染-无症状细菌尿",
    16: "男性生殖系统感染-前列腺炎", 17: "男性生殖系统感染-附睾炎", 18: "泌尿、男性生殖系统结核-泌尿系统结核",
    19: "泌尿、男性生殖系统结核-男性生殖系统结核", 20: "尿路结石-概述", 21: "尿路结石-上尿路结石",
    22: "尿路结石-膀胱结石", 23: "泌尿、男性生殖系统肿瘤-肾肿瘤（肾癌、肾母细胞瘤、肾血管平滑肌脂肪瘤）",
    24: "泌尿、男性生殖系统肿瘤-尿路上皮肿瘤（膀胱肿瘤，肾盂、输尿管癌）", 25: "泌尿、男性生殖系统肿瘤-前列腺癌",
    26: "泌尿、男性生殖系统肿瘤-睾丸肿瘤", 27: "泌尿、男性生殖系统肿瘤-阴茎癌", 28: "泌尿系统梗阻-概论",
    29: "泌尿系统梗阻-肾积水", 30: "泌尿系统梗阻-良性前列腺增生", 31: "泌尿系统梗阻-尿潴留",
    32: "泌尿系统外伤-肾外伤", 33: "泌尿系统外伤-膀胱外伤", 34: "泌尿系统外伤-前尿道外伤",
    35: "泌尿系统外伤-后尿道外伤", 36: "泌尿、男性生殖系统先天性畸形及其他疾病-隐睾",
    37: "泌尿、男性生殖系统先天性畸形及其他疾病-鞘膜积液", 38: "泌尿、男性生殖系统先天性畸形及其他疾病-精索静脉曲张",
    39: "肾功能不全-急性肾损伤（急性肾衰竭）", 40: "肾功能不全-慢性肾脏病（慢性肾衰竭）"
}


# ===== 文本格式化 =====

def format_reference(ref_str):
    """将 A1_0, A1_3, A1_1、A1_2 格式化为 A1型题第0题，A1型题第3题"""
    if not ref_str:
        return ""
    # 使用正则按照英文逗号、中文逗号、顿号、斜杠、空格进行分割
    parts = re.split(r'[,，、/ ]+', str(ref_str))
    res = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        if '_' in p:
            # 确保只取前两个部分，防止异常格式导致解包失败
            sub_parts = p.split('_')
            if len(sub_parts) >= 2:
                t, idx = sub_parts[0], sub_parts[1]
                res.append(f"{t}型题第{idx}题")
            else:
                res.append(p)
        else:
            res.append(p)
    return "，".join(res)

def parse_qgeval(s):
    """解析 QGEval 评分"""
    if not s or ':' not in s: return s
    total, scores_str = s.split(':')
    scores = [x.strip() for x in scores_str.split(',')]
    if len(scores) >= 7:
        return f"总分{total.strip()}分，流畅性{scores[0]}分，清晰度{scores[1]}分，简洁性{scores[2]}分，相关性{scores[3]}分，一致性{scores[4]}分，可回答性{scores[5]}分，答案一致性{scores[6]}分。"
    return s

def parse_ulm(s):
    """解析 ULM 评分"""
    if not s or ':' not in s: return s
    total, scores_str = s.split(':')
    scores = [x.strip() for x in scores_str.split(',')]
    if len(scores) >= 16:
        return f"总分{total.strip()}分，流畅性{scores[0]}分、排他性{scores[1]}分、明确性{scores[2]}分、目标性{scores[3]}分、综合性{scores[4]}分、侧重性{scores[5]}分、防猜性{scores[6]}分、完整性{scores[7]}分、正确性{scores[8]}分、可解性{scores[9]}分、绝对性{scores[10]}分、迷惑性{scores[11]}分、思维性{scores[12]}分、反馈性{scores[13]}分、公平性{scores[14]}分、答案解析专门评分{scores[15]}分。"
    return s


# ===== 导出流程 =====

def json_to_docx(json_path, docx_path):
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
        
    doc = Document()
    
    for item in data:
        q_type = item.get('type', '未知')
        q_id = item.get('id', '未知')
        
        # 题头分割线。
        doc.add_paragraph(f"--------------------------------------------{q_type}型题第{q_id}题--------------------------------------------")
        
        # 机器评价元数据。
        doc.add_paragraph(f"出题参考：人类题库{format_reference(item.get('prototype', ''))}。")
        doc.add_paragraph(f"FuzzyWuzzy Levenshtein：最接近人类题库{format_reference(item.get('fuzzywuzzy_doubt', ''))}，值为{item.get('fuzzywuzzy_ratio_max', '')}。")
        doc.add_paragraph(f"Sentence-BERT Cosine：最接近人类题库{format_reference(item.get('sentencebert_doubt', ''))}，值为{item.get('sentencebert_cosine_max', '')}。")
        doc.add_paragraph(f"3-gram Jaccard Index：最接近人类题库{format_reference(item.get('3gram_doubt', ''))}，值为{item.get('3gram_jaccard_max', '')}。")
        doc.add_paragraph(f"Flesch Reading Ease：{item.get('textstat_flesch_reading_ease', '')}。")
        doc.add_paragraph(f"QGEval：{parse_qgeval(item.get('QGEval', ''))}")
        doc.add_paragraph(f"ULM：{parse_ulm(item.get('ULM', ''))}")
        doc.add_paragraph("")
        
        # 考点还原字段可能包含一个或多个编号。
        tp_list = item.get('test_point', [])
        tp_str = "、".join([TEST_POINTS.get(tp, str(tp)) for tp in tp_list])
        
        if q_type in ['A1', 'A2', 'X']:
            # 单问题型直接写题干、选项、考点、答案和解析。
            doc.add_paragraph(item.get('stem', ''))
            opts = item.get('options', {})
            for k, v in opts.items():
                doc.add_paragraph(f"{k}. {v}")
            doc.add_paragraph(f"考点还原：{tp_str}。")
            doc.add_paragraph(f"参考答案：{item.get('answer', '')}。")
            doc.add_paragraph(f"答案解析：{item.get('analysis', '')}")
            
        else:
            # 多问题型按 stem1/stem2... 写入各小题。
            sub_count = sum(1 for k in item.keys() if k.startswith('stem') and k != 'stem')
            
            for i in range(1, sub_count + 1):
                stem = item.get(f'stem{i}', '')
                
                # A3/A4 的病例材料并入第一小题。
                if i == 1 and 'case' in item:
                    stem = item['case'] + stem
                    
                doc.add_paragraph(f"({i}) {stem}")
                
                # 子问题优先使用独立选项；B 型题可共用 options。
                opts = item.get(f'options{i}', item.get('options', {}))
                for k, v in opts.items():
                    doc.add_paragraph(f"{k}. {v}")
                    
                doc.add_paragraph(f"考点还原：{tp_str}。")
                doc.add_paragraph(f"参考答案：{item.get(f'answer{i}', '')}。")
                doc.add_paragraph(f"答案解析：{item.get(f'analysis{i}', '')}")
                
        # 每道大题之后保留空行。
        doc.add_paragraph("")

    doc.save(docx_path)
    print(f"成功导出：{docx_path}")


# ===== 程序入口 =====

if __name__ == "__main__":
    # 批量处理 data/banks 下的 MAS 题库 JSON，输出为给合作者人工审阅的 Word 文档。
    files_to_convert = [
        "new_bank_a1.json",
        "new_bank_a2.json",
        "new_bank_a3.json",
        "new_bank_a4.json",
        "new_bank_b.json",
        "new_bank_x.json"
    ]
    
    for file in files_to_convert:
        json_path = BANK_DIR / file
        if json_path.exists():
            REPORT_EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
            out_name = REPORT_EXPORTS_DIR / file.replace('.json', '.docx')
            json_to_docx(str(json_path), str(out_name))
        else:
            print(f"未找到文件: {json_path}")
